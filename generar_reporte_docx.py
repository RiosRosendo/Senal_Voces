"""
Genera el reporte en formato .docx con imágenes incrustadas.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Márgenes ──────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Estilos base ──────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading(paragraph, text, level=1):
    paragraph.text = text
    paragraph.style = f'Heading {level}'

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_paragraph(doc, text='', bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p

def add_code(doc, code_text):
    """Agrega un bloque de código con fondo gris."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # fondo gris claro via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    return p

def add_image(doc, path, caption, width=Inches(5.5)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    else:
        p = doc.add_paragraph(f'[Imagen no encontrada: {path}]')
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Encabezados
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'D6E4F0')
        tcPr.append(shd)
    # Filas
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = str(val)
            row.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


# ═══════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Reconocimiento de Voz\ncon Cuantización Vectorial')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

inst = doc.add_paragraph()
inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = inst.add_run('Instituto Tecnológico y de Estudios Superiores de Monterrey')
r.font.size = Pt(13)
r.bold = True

mat = doc.add_paragraph()
mat.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = mat.add_run('Procesamiento Digital de Señales')
r2.font.size = Pt(12)
r2.italic = True

doc.add_paragraph()
doc.add_paragraph()

info = [
    ('Integrantes', 'Rosendo De Los Ríos Moreno    —    A01198515'),
    ('',            'Juan José Jáuregui Barba       —    A00836722'),
    ('Fecha',       'Abril 2026'),
]
tbl = doc.add_table(rows=len(info), cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (label, value) in enumerate(info):
    tbl.rows[ri].cells[0].text = label
    tbl.rows[ri].cells[0].paragraphs[0].runs[0].bold = True
    tbl.rows[ri].cells[1].text = value
for row in tbl.rows:
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 1. OBJETIVO
# ═══════════════════════════════════════════════════
add_heading(doc, '1. Objetivo')
add_paragraph(doc,
    'El objetivo de esta práctica fue diseñar e implementar un sistema básico de '
    'reconocimiento de voz utilizando Cuantización Vectorial (VQ). El sistema '
    'aprende a distinguir 10 palabras en inglés a partir de grabaciones de voz '
    'y las clasifica usando modelos estadísticos basados en características espectrales (LPC/LSF).')
add_paragraph(doc,
    'Esta actividad forma parte de un desarrollo progresivo hacia la integración '
    'de comandos de voz en un robot móvil (Puzzlebot con ROS2).', italic=True)

# ═══════════════════════════════════════════════════
# 2. PALABRAS
# ═══════════════════════════════════════════════════
add_heading(doc, '2. Palabras Utilizadas')
add_paragraph(doc, 'Se trabajó con 10 palabras de comandos de navegación robótica:')
add_table(doc,
    ['#', 'Palabra', '#', 'Palabra'],
    [['1','start','6','backward'],
     ['2','stop','7','lift'],
     ['3','left','8','leave'],
     ['4','right','9','break'],
     ['5','forward','10','continue']])
doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 3. PIPELINE
# ═══════════════════════════════════════════════════
add_heading(doc, '3. Pipeline del Sistema')
add_paragraph(doc, 'El sistema siguió un flujo de procesamiento de 6 etapas:')
add_code(doc,
"""Grabación de voz (iPhone, formato M4A)
      ↓
Segmentación automática por energía (0_segmentar_audio.py)
      ↓
Filtro de Preénfasis  H(z) = 1 − 0.95·z⁻¹
      ↓
Ventaneo Hamming (320 muestras, salto 128)
      ↓
Detección de actividad de voz (VAD)
      ↓
Extracción LPC orden 12  →  Conversión a LSF
      ↓
Cuantización Vectorial K-means (K = 16, 32, 64)
      ↓
Clasificación con distancia Itakura-Saito
      ↓
Matriz de Confusión""")

# ═══════════════════════════════════════════════════
# 4. ADQUISICIÓN
# ═══════════════════════════════════════════════════
add_heading(doc, '4. Adquisición de Datos')
add_heading(doc, '4.1 Método de Grabación', level=2)
add_paragraph(doc,
    'Debido a problemas con el micrófono interno del equipo de cómputo (ruido eléctrico '
    'y DC bias), las grabaciones se realizaron con un iPhone usando la aplicación de '
    'grabadora de voz nativa. Los archivos fueron exportados en formato .m4a.')
add_paragraph(doc, 'Protocolo de grabación:', bold=True)
for item in ['Un archivo de audio por palabra (10 archivos en total)',
             'Cada archivo contiene la palabra repetida 15 veces con ~1 segundo de pausa',
             'Ambiente silencioso y volumen de voz consistente',
             'Misma persona en todas las grabaciones']:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(11)

add_heading(doc, '4.2 Segmentación Automática', level=2)
add_paragraph(doc,
    'El script 0_segmentar_audio.py procesó cada audio y lo dividió en 15 archivos '
    'individuales usando detección por umbral de energía:')
add_code(doc,
"""# Umbral de energía para detectar voz
e_thresh = 0.015 * np.max(energy)
voiced   = energy > e_thresh

# Fusionar regiones separadas por silencio corto (<0.4 s)
max_gap = int(0.4 * fs / hop)
for r in regiones:
    if fusionadas and r[0] - fusionadas[-1][1] <= max_gap:
        fusionadas[-1][1] = r[1]   # fusionar
    else:
        fusionadas.append(r)       # nueva región""")

doc.add_paragraph()
add_image(doc,
    'resultados/segmentacion/start_segmentacion.png',
    'Figura 1 — Segmentación de "start": 15 repeticiones detectadas (regiones verdes)')
doc.add_paragraph()
add_image(doc,
    'resultados/segmentacion/continue_segmentacion.png',
    'Figura 2 — Segmentación de "continue": palabra más larga del vocabulario')
doc.add_paragraph()

add_heading(doc, '4.3 Estadísticas de Grabaciones', level=2)
add_table(doc,
    ['Palabra','Archivos','Entrenamiento','Prueba'],
    [['start','15','10','5'],
     ['stop','15','10','5'],
     ['left','15','10','5'],
     ['right','15','10','5'],
     ['forward','15','10','5'],
     ['backward','15','10','5'],
     ['lift','15','10','5'],
     ['leave','15','10','5'],
     ['break','15','10','5'],
     ['continue','15','10','5'],
     ['TOTAL','150','100','50']])
doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 5. PREPROCESAMIENTO
# ═══════════════════════════════════════════════════
add_heading(doc, '5. Preprocesamiento')

add_heading(doc, '5.1 Filtro de Preénfasis', level=2)
add_paragraph(doc,
    'Se aplicó el filtro de preénfasis para amplificar las frecuencias altas de '
    'la señal de voz, compensando la caída natural del espectro vocal:')
add_paragraph(doc, '        H(z) = 1 − 0.95 · z⁻¹        →        y[n] = x[n] − 0.95·x[n−1]',
              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_code(doc,
"""def apply_preemphasis(signal, alpha=0.95):
    return np.concatenate([[signal[0]],
                           signal[1:] - alpha * signal[:-1]])""")

add_heading(doc, '5.2 Ventaneo Hamming', level=2)
add_paragraph(doc,
    'La señal se dividió en frames superpuestos con ventana de Hamming:')
add_table(doc,
    ['Parámetro','Valor','Equivalente temporal'],
    [['Longitud de ventana','320 muestras','20 ms'],
     ['Desplazamiento (hop)','128 muestras','8 ms'],
     ['Solapamiento','~60%','—']])
doc.add_paragraph()
add_code(doc,
"""def hamming_frames(signal, frame_len=320, hop=128):
    n_frames = 1 + (len(signal) - frame_len) // hop
    window   = np.hamming(frame_len)
    frames   = np.zeros((n_frames, frame_len))
    for i in range(n_frames):
        start = i * hop
        frames[i] = signal[start:start+frame_len] * window
    return frames""")

add_heading(doc, '5.3 Detección de Actividad de Voz (VAD)', level=2)
add_paragraph(doc,
    'Se usó un umbral adaptativo de energía para detectar el inicio y fin '
    'de cada palabra dentro del audio. El umbral se establece como el 2% '
    'de la energía máxima del archivo:')
add_code(doc,
"""e_thresh = 0.02 * np.max(energy)
voice    = energy > e_thresh
# Margen de 4 frames al inicio y fin para no cortar la palabra
first_frame = max(0, voiced_frames[0]  - 4)
last_frame  = min(n_frames-1, voiced_frames[-1] + 4)""")
doc.add_paragraph()
add_image(doc,
    'resultados/preprocesamiento/backward/01.png',
    'Figura 3 — Preprocesamiento de "backward": señal con VAD (arriba) y señal recortada (abajo)')
doc.add_paragraph()
add_image(doc,
    'resultados/preprocesamiento/forward/05.png',
    'Figura 4 — Preprocesamiento de "forward": detección clara de inicio y fin de palabra')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 6. EXTRACCIÓN DE CARACTERÍSTICAS
# ═══════════════════════════════════════════════════
add_heading(doc, '6. Extracción de Características')

add_heading(doc, '6.1 Codificación Predictiva Lineal (LPC)', level=2)
add_paragraph(doc,
    'Para cada frame se calcularon 12 coeficientes LPC usando el método de '
    'autocorrelación con el algoritmo de Levinson-Durbin. El LPC modela la '
    'señal de voz como la salida de un filtro todo-polo excitado por ruido blanco:')
add_code(doc,
"""def compute_lpc(frame, order=12):
    # Autocorrelación del frame
    r = np.array([np.dot(frame[:n-k], frame[k:])
                  for k in range(order + 1)])
    # Levinson-Durbin
    a, e = np.zeros(order), r[0]
    for m in range(order):
        lam = -(r[m+1] + np.dot(a[:m], r[m:0:-1])) / e
        a_new = a.copy()
        if m > 0:
            a_new[:m] = a[:m] + lam * a[m-1::-1]
        a_new[m] = lam
        a, e = a_new, e * (1 - lam**2)
    return a, e    # coeficientes LPC y ganancia de predicción""")

add_heading(doc, '6.2 Frecuencias Espectrales de Línea (LSF)', level=2)
add_paragraph(doc,
    'Los coeficientes LPC se convirtieron a LSF (Line Spectral Frequencies) '
    'para el proceso de clustering. Las LSF tienen mejor comportamiento numérico '
    'y representan directamente las resonancias del tracto vocal en el rango (0, π):')
add_code(doc,
"""def lpc_to_lsf(lpc):
    a     = np.concatenate([[1.0], lpc])
    P     = a_pad + a_rev_pad    # polinomio simétrico
    Q     = a_pad - a_rev_pad    # polinomio antisimétrico
    # Raíces sobre el círculo unitario → ángulos = LSF
    P_red, _ = np.polydiv(P, [1.0,  1.0])   # quitar raíz z = -1
    Q_red, _ = np.polydiv(Q, [1.0, -1.0])   # quitar raíz z = +1
    # Extraer ángulos positivos de las raíces
    return np.array(sorted(p_angles + q_angles)[:order])""")

add_heading(doc, '6.3 Frames por Palabra', level=2)
add_table(doc,
    ['Palabra','Frames/grabación (promedio)','Total frames entrenamiento'],
    [['start','64','640'],
     ['stop','52','522'],
     ['left','45','447'],
     ['right','43','434'],
     ['forward','63','632'],
     ['backward','56','558'],
     ['lift','47','469'],
     ['leave','44','437'],
     ['break','40','398'],
     ['continue','81','806']])
doc.add_paragraph()
add_image(doc,
    'resultados/caracteristicas/start/01.png',
    'Figura 5 — Evolución de los 12 coeficientes LSF frame a frame para "start"')
doc.add_paragraph()
add_image(doc,
    'resultados/caracteristicas/continue/01.png',
    'Figura 6 — Coeficientes LSF para "continue" (mayor duración = más frames)')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 7. CUANTIZACIÓN VECTORIAL
# ═══════════════════════════════════════════════════
add_heading(doc, '7. Cuantización Vectorial (Entrenamiento)')
add_paragraph(doc,
    'La Cuantización Vectorial crea un codebook (libro de códigos) por cada palabra. '
    'El codebook es un conjunto de vectores representativos (centroides) que resumen '
    'las características espectrales de esa palabra usando K-means:')
add_code(doc,
"""from sklearn.cluster import KMeans

for K in [16, 32, 64]:
    kmeans = KMeans(n_clusters=K, n_init=10,
                    max_iter=300, random_state=42)
    kmeans.fit(lsf_data)              # (N_frames_total, 12)
    centroids = kmeans.cluster_centers_   # (K, 12)
    np.savez_compressed(f'{palabra}_K{K}.npz', centroids=centroids)""")
doc.add_paragraph()
add_table(doc,
    ['Tamaño K','Centroides por palabra','Total codebooks generados'],
    [['16','16 vectores de 12 dimensiones','10'],
     ['32','32 vectores de 12 dimensiones','10'],
     ['64','64 vectores de 12 dimensiones','10']])
doc.add_paragraph()
add_image(doc,
    'resultados/codebooks/start_centroides.png',
    'Figura 7 — Centroides del codebook de "start" para K=16, 32 y 64')
doc.add_paragraph()
add_image(doc,
    'resultados/codebooks/backward_centroides.png',
    'Figura 8 — Centroides del codebook de "backward"')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 8. DISTANCIA ITAKURA-SAITO
# ═══════════════════════════════════════════════════
add_heading(doc, '8. Clasificación con Distancia de Itakura-Saito')
add_paragraph(doc,
    'Para clasificar una palabra de prueba se calcula la distorsión VQ usando '
    'la distancia de Itakura-Saito entre cada frame del audio y los vectores del codebook. '
    'Esta distancia mide cuánto peor predice el modelo B la señal del modelo A:')
add_paragraph(doc,
    '    d_IS(A→B) = (b·R_A·b) / σ²_A  −  log((b·R_A·b) / σ²_A)  −  1',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc,
    'Donde R_A es la matriz de autocorrelación del frame A, b es el vector LPC '
    'del centroide B, y σ²_A es la ganancia de predicción del frame A.')
add_code(doc,
"""def itakura_saito_distance(lpc_a, gain_a, acf_a, lpc_b):
    b_full = np.concatenate([[1.0], lpc_b])
    R      = toeplitz(acf_a[:order + 1])
    num    = float(b_full @ R @ b_full)
    ratio  = num / (gain_a + 1e-12)
    dist   = ratio - np.log(ratio) - 1.0
    return max(0.0, dist)""")
add_paragraph(doc,
    'La palabra se clasifica como aquella cuyo codebook produce la menor '
    'distorsión promedio sobre todos los frames del audio de prueba.')

# ═══════════════════════════════════════════════════
# 9. EVALUACIÓN
# ═══════════════════════════════════════════════════
add_heading(doc, '9. Evaluación del Sistema')
add_paragraph(doc,
    'Se usaron las 5 grabaciones restantes (archivos 11 al 15) de cada palabra '
    'como conjunto de prueba — 50 muestras en total.')
doc.add_paragraph()
add_image(doc,
    'resultados/evaluacion/confusion_K16.png',
    'Figura 9 — Matriz de Confusión con K = 16')
doc.add_paragraph()
add_image(doc,
    'resultados/evaluacion/confusion_K32.png',
    'Figura 10 — Matriz de Confusión con K = 32')
doc.add_paragraph()
add_image(doc,
    'resultados/evaluacion/confusion_K64.png',
    'Figura 11 — Matriz de Confusión con K = 64')
doc.add_paragraph()
add_image(doc,
    'resultados/evaluacion/comparativa_K.png',
    'Figura 12 — Comparativa de accuracy global por tamaño de codebook')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 10. ANÁLISIS
# ═══════════════════════════════════════════════════
add_heading(doc, '10. Análisis de Resultados')

add_heading(doc, '10.1 Efecto del Tamaño del Codebook', level=2)
add_table(doc,
    ['K','Descripción'],
    [['16','Codebook pequeño. Rápido pero puede subrepresentar variaciones en la pronunciación'],
     ['32','Codebook mediano. Balance entre memoria y representación espectral'],
     ['64','Codebook grande. Mayor capacidad representativa, requiere más datos de entrenamiento']])
doc.add_paragraph()

add_heading(doc, '10.2 Palabras con Mayor Confusión', level=2)
add_paragraph(doc,
    'Las palabras con más errores de clasificación tienden a ser fonéticamente similares:')
for item in ['"left" vs "lift" — comparten el fonema /l/ y terminan en /t/',
             '"stop" vs "start" — mismo inicio consonántico /st/',
             '"leave" vs "lift" — inicio /l/ similar']:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(11)

add_heading(doc, '10.3 Limitaciones del Sistema', level=2)
for item in ['Solo un locutor: no generaliza a otras voces',
             'Sin normalización de duración de la señal',
             'Sensible al ruido de fondo',
             'Solo 10 muestras de entrenamiento por palabra']:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(11)

# ═══════════════════════════════════════════════════
# 11. CONCLUSIONES
# ═══════════════════════════════════════════════════
add_heading(doc, '11. Conclusiones')
conclusiones = [
    'El sistema LPC + LSF + VQ es funcional para vocabularios pequeños con un solo locutor.',
    'La distancia de Itakura-Saito es apropiada para comparar modelos espectrales LPC.',
    'El tamaño del codebook impacta el rendimiento: K más grande representa mejor '
    'la variabilidad vocal pero requiere más datos de entrenamiento.',
    'La calidad de las grabaciones es el principal factor de éxito del sistema. '
    'Un protocolo estricto (mismo locutor, ambiente silencioso) mejora el accuracy.',
    'Este sistema sienta las bases para integración futura con ROS2 como nodo '
    'publisher de comandos de voz para el robot Puzzlebot.'
]
for i, c in enumerate(conclusiones, 1):
    p = doc.add_paragraph(f'{i}. {c}')
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════
# APÉNDICE
# ═══════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, 'Apéndice — Parámetros del Sistema')
add_table(doc,
    ['Parámetro','Valor','Justificación'],
    [['Frecuencia de muestreo','16,000 Hz','Estándar para voz, captura formantes hasta 8 kHz'],
     ['Coeficiente preénfasis α','0.95','Compensación estándar del espectro vocal'],
     ['Longitud de ventana','320 muestras (20 ms)','Resolución temporal adecuada para fonemas'],
     ['Desplazamiento (hop)','128 muestras (8 ms)','Solapamiento del 60%'],
     ['Orden LPC','12','Captura los formantes principales de la voz'],
     ['Tamaños codebook K','16, 32, 64','Comparación de capacidad representativa'],
     ['Muestras entrenamiento','10 por palabra','67% del total (150 grabaciones)'],
     ['Muestras prueba','5 por palabra','33% del total']])

doc.add_paragraph()
p = doc.add_paragraph(
    'Reporte generado para la asignatura de Procesamiento Digital de Señales — ITESM — Abril 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True
p.runs[0].font.size = Pt(9)

# ── Guardar ───────────────────────────────────────
out = 'Reporte_Reconocimiento_Voz_VQ.docx'
doc.save(out)
print(f'✓ Reporte guardado: {out}')
